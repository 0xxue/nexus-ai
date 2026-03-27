"""
Document Parser — Extract text from PDF, Word, Excel files

Supports:
  - PDF: pdfplumber (tables + text)
  - Word (.docx): python-docx
  - Excel (.xlsx/.csv): pandas
  - Text (.txt/.md): direct read

Each parser returns a list of chunks with metadata (page, section, etc.)
"""

import os
import structlog
from typing import Optional
from dataclasses import dataclass

logger = structlog.get_logger()


@dataclass
class ParsedChunk:
    """A chunk of extracted text with metadata."""
    content: str
    metadata: dict  # page, section, row_range, etc.


class DocumentParser:
    """Parse documents into text chunks."""

    SUPPORTED_TYPES = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".csv", ".txt", ".md"}

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def parse(self, file_path: str) -> list[ParsedChunk]:
        """Auto-detect file type and parse."""
        ext = os.path.splitext(file_path)[1].lower()

        if ext not in self.SUPPORTED_TYPES:
            raise ValueError(f"Unsupported file type: {ext}")

        logger.info("Parsing document", path=file_path, type=ext)

        match ext:
            case ".pdf":
                return self._parse_pdf(file_path)
            case ".docx" | ".doc":
                return self._parse_word(file_path)
            case ".xlsx" | ".xls":
                return self._parse_excel(file_path)
            case ".csv":
                return self._parse_csv(file_path)
            case ".txt" | ".md":
                return self._parse_text(file_path)
            case _:
                return self._parse_text(file_path)

    def _parse_pdf(self, path: str) -> list[ParsedChunk]:
        """Extract text from PDF page by page, then chunk."""
        try:
            import pdfplumber
        except ImportError:
            logger.warning("pdfplumber not installed, trying fallback")
            return self._parse_text(path)

        chunks = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    page_chunks = self._split_text(text, metadata={"page": i + 1, "source": "pdf"})
                    chunks.extend(page_chunks)

                # Extract tables as separate chunks
                tables = page.extract_tables()
                for j, table in enumerate(tables):
                    table_text = self._table_to_text(table)
                    if table_text:
                        chunks.append(ParsedChunk(
                            content=table_text,
                            metadata={"page": i + 1, "type": "table", "table_index": j},
                        ))

        logger.info("PDF parsed", pages=len(pdf.pages) if pdf else 0, chunks=len(chunks))
        return chunks

    def _parse_word(self, path: str) -> list[ParsedChunk]:
        """Extract text from Word document."""
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx not installed")
            return []

        doc = Document(path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        combined = "\n".join(full_text)
        chunks = self._split_text(combined, metadata={"source": "docx"})

        # Extract tables
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_text = "\n".join(rows)
            if table_text:
                chunks.append(ParsedChunk(
                    content=table_text,
                    metadata={"type": "table", "table_index": i, "source": "docx"},
                ))

        logger.info("Word parsed", paragraphs=len(doc.paragraphs), chunks=len(chunks))
        return chunks

    def _parse_excel(self, path: str) -> list[ParsedChunk]:
        """Extract data from Excel, each sheet as chunks."""
        try:
            import pandas as pd
        except ImportError:
            logger.warning("pandas not installed")
            return []

        chunks = []
        xl = pd.ExcelFile(path)

        for sheet_name in xl.sheet_names:
            df = pd.read_excel(xl, sheet_name=sheet_name)
            if df.empty:
                continue

            # Convert to text representation
            text = f"Sheet: {sheet_name}\n"
            text += f"Columns: {', '.join(df.columns.astype(str))}\n"
            text += f"Rows: {len(df)}\n\n"

            # Chunk rows in groups
            row_size = 20
            for start in range(0, len(df), row_size):
                end = min(start + row_size, len(df))
                subset = df.iloc[start:end]
                chunk_text = text + subset.to_string(index=False)
                chunks.append(ParsedChunk(
                    content=chunk_text,
                    metadata={"sheet": sheet_name, "rows": f"{start + 1}-{end}", "source": "excel"},
                ))

        logger.info("Excel parsed", sheets=len(xl.sheet_names), chunks=len(chunks))
        return chunks

    def _parse_csv(self, path: str) -> list[ParsedChunk]:
        """Parse CSV file."""
        try:
            import pandas as pd
        except ImportError:
            return self._parse_text(path)

        df = pd.read_csv(path)
        chunks = []
        row_size = 20

        for start in range(0, len(df), row_size):
            end = min(start + row_size, len(df))
            subset = df.iloc[start:end]
            chunks.append(ParsedChunk(
                content=subset.to_string(index=False),
                metadata={"rows": f"{start + 1}-{end}", "source": "csv"},
            ))

        return chunks

    def _parse_text(self, path: str) -> list[ParsedChunk]:
        """Parse plain text / markdown."""
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
        return self._split_text(text, metadata={"source": "text"})

    def _split_text(self, text: str, metadata: dict = None) -> list[ParsedChunk]:
        """
        Split text into chunks with overlap.
        Tries to split on paragraph boundaries first, falls back to character split.
        """
        if not text.strip():
            return []

        metadata = metadata or {}
        chunks = []

        # Split by paragraphs first
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

        current_chunk = ""
        for para in paragraphs:
            if len(current_chunk) + len(para) <= self.chunk_size:
                current_chunk += para + "\n\n"
            else:
                if current_chunk:
                    chunks.append(ParsedChunk(
                        content=current_chunk.strip(),
                        metadata={**metadata, "chunk_index": len(chunks)},
                    ))
                current_chunk = para + "\n\n"

        if current_chunk.strip():
            chunks.append(ParsedChunk(
                content=current_chunk.strip(),
                metadata={**metadata, "chunk_index": len(chunks)},
            ))

        # If no paragraphs found, do character-level splitting
        if not chunks and text.strip():
            for i in range(0, len(text), self.chunk_size - self.chunk_overlap):
                chunk = text[i:i + self.chunk_size]
                if chunk.strip():
                    chunks.append(ParsedChunk(
                        content=chunk.strip(),
                        metadata={**metadata, "chunk_index": len(chunks), "char_offset": i},
                    ))

        return chunks

    @staticmethod
    def _table_to_text(table: list) -> str:
        """Convert a table (list of rows) to text."""
        if not table:
            return ""
        rows = []
        for row in table:
            cells = [str(cell).strip() if cell else "" for cell in row]
            rows.append(" | ".join(cells))
        return "\n".join(rows)
